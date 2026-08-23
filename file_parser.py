"""File parsing for BPCommentary bot: PDF, Word, and image understanding."""

from __future__ import annotations

import base64
import io
import logging
from typing import Optional

import httpx

logger = logging.getLogger(__name__)

# Vision model for product images
VISION_MODEL = "meta-llama/Llama-3.2-90B-Vision-Instruct-Turbo"
VISION_MAX_TOKENS = 2048

# Supported file types
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"
IMAGE_MIMES = {"image/jpeg", "image/png", "image/webp", "image/gif"}

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB Telegram bot limit
MAX_TEXT_CHARS = 15000  # cap extracted text to avoid LLM token bloat


async def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "[PDF extraction requires PyPDF2 library]"

    text_parts: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text.strip())
    except Exception as e:
        logger.exception("PDF extraction failed")
        return f"[PDF parsing error: {e}]"

    text = "\n\n".join(text_parts).strip()
    if not text:
        return "[PDF contains no extractable text — it may be scanned images. Try sending the text directly.]"
    return text[:MAX_TEXT_CHARS]


async def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx Word file."""
    try:
        import docx
    except ImportError:
        return "[Word extraction requires python-docx library]"

    text_parts: list[str] = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
    except Exception as e:
        logger.exception("DOCX extraction failed")
        return f"[Word parsing error: {e}]"

    text = "\n".join(text_parts).strip()
    if not text:
        return "[Word document contains no extractable text.]"
    return text[:MAX_TEXT_CHARS]


async def analyze_image(file_bytes: bytes, mime_type: str, caption: str = "") -> str:
    """Send an image to the vision LLM and get a business-plan-relevant description."""
    import os

    api_key = os.getenv("TOGETHER_API_KEY", "")
    if not api_key:
        return "[Image analysis requires TOGETHER_API_KEY]"

    b64 = base64.b64encode(file_bytes).decode("ascii")
    data_url = f"data:{mime_type};base64,{b64}"

    prompt = (
        "You are analyzing a product image or business plan visual for a startup audit. "
        "Describe in detail what you see: product design, branding, packaging, materials, "
        "target market signals, price positioning clues, and any text visible in the image. "
        "If this is a product photo, assess the aesthetic quality, luxury feel, and market positioning. "
        "Be concrete and specific. Do not speculate beyond what is visible."
    )
    if caption:
        prompt += f"\n\nUser's note/caption: {caption}"

    payload = {
        "model": VISION_MODEL,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        "max_tokens": VISION_MAX_TOKENS,
        "temperature": 0.3,
    }

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.together.xyz/v1/chat/completions",
                json=payload,
                headers={"Authorization": f"Bearer {api_key}"},
            )
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()[:MAX_TEXT_CHARS]
    except Exception as e:
        logger.exception("Vision API failed")
        return f"[Image analysis failed: {e}]"


async def parse_telegram_file(update, context, caption: str = "") -> Optional[str]:
    """
    Download and parse a file sent to the bot (document or photo).
    Returns extracted text, or None if the file type is unsupported.
    """
    msg = update.message
    if msg is None:
        return None

    # Determine file: document or photo
    if msg.document:
        doc = msg.document
        mime = doc.mime_type or ""
        file_id = doc.file_id
        file_name = doc.file_name or "file"
    elif msg.photo:
        # Telegram sends multiple sizes; pick the largest
        photo = msg.photo[-1]
        mime = "image/jpeg"
        file_id = photo.file_id
        file_name = "photo.jpg"
    else:
        return None

    # Check size
    file_obj = await context.bot.get_file(file_id)
    if file_obj.file_size and file_obj.file_size > MAX_FILE_BYTES:
        return f"[File too large: {file_obj.file_size // 1024 // 1024}MB. Max 20MB.]"

    # Download into memory
    buf = io.BytesIO()
    await file_obj.download_to_memory(buf)
    file_bytes = buf.getvalue()

    # Route by type
    if msg.document and mime == PDF_MIME:
        text = await extract_pdf(file_bytes)
        return f"[PDF: {file_name}]\n\n{text}"
    elif msg.document and mime in (DOCX_MIME, DOC_MIME):
        text = await extract_docx(file_bytes)
        return f"[Word: {file_name}]\n\n{text}"
    elif mime in IMAGE_MIMES or msg.photo:
        text = await analyze_image(file_bytes, mime, caption)
        return f"[Image: {file_name}]\n\n{text}"
    else:
        return None
